from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw
from reportlab.lib.colors import HexColor
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer


OUTPUT_DIR = Path(__file__).resolve().parent
BLUE = "036EB7"
INK = "17323D"
MUTED = "526A75"


def set_font(run, name: str, size: float, color: str, bold: bool = False):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold


def build_docx():
    document = Document()
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for style_name, size, before, after in [
        ("Heading 1", 16, 16, 8),
        ("Heading 2", 13, 12, 6),
    ]:
        style = document.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(BLUE)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    title = document.add_paragraph()
    title.paragraph_format.space_after = Pt(4)
    set_font(title.add_run("Field Observation Follow-up Note"), "Calibri", 22, INK, True)
    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    set_font(subtitle.add_run("De-identified test fixture | Harbor Community Day Center | 2026-08-18"), "Calibri", 10, MUTED)

    document.add_heading("Observed signal", level=1)
    document.add_paragraph(
        "Across three afternoon activity sessions, staff recorded shorter periods of sustained attention, "
        "more requests to leave early, and lower participation during repeated worksheet-based activities. "
        "The same participants remained engaged longer during live music and small-group storytelling."
    )

    document.add_heading("Psychological concern to explore", level=1)
    document.add_paragraph(
        "The pattern may indicate reduced engagement and a need for stronger social connection. Two residents "
        "also initiated repeated conversations about having fewer family visits. These observations support "
        "a loneliness-related concern for follow-up, but they do not establish a diagnosis."
    )

    document.add_heading("Uncertainty", level=1)
    document.add_paragraph(
        "The current evidence cannot distinguish activity repetition, session length, fatigue, hearing access, "
        "or a sustained cognitive or emotional change. Baseline attention, activity duration, and response after "
        "changing the format were not recorded consistently."
    )

    document.add_heading("Next collection prompts", level=1)
    document.add_paragraph(
        "At the next visit, record the activity type, the approximate time when attention changes, whether the "
        "pattern repeats for the same person, and whether engagement improves after switching to a social or "
        "movement-based format. Ask a neutral follow-up about desired companionship and preferred contact."
    )

    footer = section.footer.paragraphs[0]
    footer.alignment = 2
    set_font(footer.add_run("CNPAF Community | Synthetic, de-identified QA evidence"), "Calibri", 8, MUTED)
    document.save(OUTPUT_DIR / "field-observation-follow-up.docx")


def build_pdf():
    styles = getSampleStyleSheet()
    body = ParagraphStyle(
        "Body",
        parent=styles["BodyText"],
        fontName="Helvetica",
        fontSize=10.5,
        leading=14,
        textColor=HexColor(f"#{INK}"),
        spaceAfter=8,
    )
    heading = ParagraphStyle(
        "Heading",
        parent=body,
        fontName="Helvetica-Bold",
        fontSize=13,
        leading=16,
        textColor=HexColor(f"#{BLUE}"),
        spaceBefore=10,
        spaceAfter=5,
    )
    title = ParagraphStyle(
        "FixtureTitle",
        parent=body,
        fontName="Helvetica-Bold",
        fontSize=20,
        leading=24,
        textColor=HexColor(f"#{INK}"),
        spaceAfter=4,
    )
    meta = ParagraphStyle(
        "Meta",
        parent=body,
        fontSize=9,
        textColor=HexColor(f"#{MUTED}"),
        spaceAfter=14,
    )
    pdf = SimpleDocTemplate(
        str(OUTPUT_DIR / "activity-engagement-summary.pdf"),
        pagesize=letter,
        leftMargin=inch,
        rightMargin=inch,
        topMargin=inch,
        bottomMargin=inch,
        title="Activity Engagement Summary",
        author="CNPAF Community QA",
    )
    story = [
        Paragraph("Activity Engagement Summary", title),
        Paragraph("Synthetic evidence attachment | Four observation sessions | No personal identifiers", meta),
        Paragraph("Pattern", heading),
        Paragraph(
            "Attention declined most often after 20 minutes in repetitive seated activities. Engagement was "
            "more stable during familiar music, paired conversation, and short movement breaks. Early departure "
            "was recorded in three of four worksheet sessions and in none of the music sessions.", body),
        Paragraph("Concern boundary", heading),
        Paragraph(
            "The pattern raises a participation and social-connection concern. It should not be described as "
            "cognitive decline without repeated person-level observations, baseline information, and alternative "
            "explanations such as fatigue, hearing access, session timing, or activity design.", body),
        Paragraph("Evidence gap", heading),
        Paragraph(
            "The source records do not consistently capture session duration, room noise, hearing-device use, "
            "individual baseline attention, or whether a format change improved participation.", body),
        Paragraph("Recommended verification", heading),
        Paragraph(
            "Compare participation duration across activity types for two weeks. Record the first sign of "
            "disengagement, the context, any accommodation offered, and the response after the activity changes.", body),
        Spacer(1, 12),
        Paragraph("Prepared only for end-to-end product testing.", meta),
    ]
    pdf.build(story)


def build_image():
    image = Image.new("RGB", (1400, 900), "white")
    draw = ImageDraw.Draw(image)
    draw.rounded_rectangle((70, 70, 1330, 830), radius=24, outline=f"#{BLUE}", width=6)
    draw.text((120, 125), "Activity session observation board", fill=f"#{INK}")
    draw.text((120, 210), "Session A: worksheet | attention change at 22 min", fill=f"#{MUTED}")
    draw.text((120, 285), "Session B: live music | sustained participation", fill=f"#{MUTED}")
    draw.text((120, 360), "Session C: storytelling | conversation increased", fill=f"#{MUTED}")
    draw.text((120, 470), "Synthetic QA image - no identifiable people", fill=f"#{BLUE}")
    image.save(OUTPUT_DIR / "activity-session-observation.png", optimize=True)


if __name__ == "__main__":
    build_docx()
    build_pdf()
    build_image()
